"""
redactor.py
-----------
End-to-end PII redaction for .docx files.

Usage:
    python3 redactor.py <input.docx> <output.docx> [--audit audit.csv]

Approach
--------
Two passes over the document:

  PASS 1 (gazetteer build): scan every paragraph (body, tables, headers,
  footers) with the high-precision, context-anchored detectors (role-anchored
  person names, company-suffix names). This builds a small, high-confidence
  gazetteer of real entities that appear in this specific document.

  PASS 2 (redaction): scan every paragraph again with the FULL detector set,
  now also matching the gazetteer entries literally wherever they appear
  (not just next to their original anchoring context), plus all
  structured detectors (email, phone, IP, SSN, credit card, DOB, address,
  CIN). Overlapping spans are resolved by category priority + longest match.
  Each paragraph's runs are merged into one run carrying the fully redacted
  text (this keeps the code simple; it can lose intra-paragraph run-level
  formatting such as a mid-sentence bold word -- see README trade-offs).

Extending to a new PII type: add a `find_x()` function to pii_detectors.py,
register it in DETECTORS, add a fake_x() to fake_values.py and map it in
FakeValueFactory.get(). No changes needed here.
"""

import csv
import re
import sys

from docx import Document

import pii_detectors as det
from fake_values import FakeValueFactory

# Priority used to resolve overlapping spans: lower number wins.
CATEGORY_PRIORITY = {
    "EMAIL": 0, "IP_ADDRESS": 0, "SSN": 0, "CREDIT_CARD": 0, "COMPANY_ID": 0,
    "DATE_OF_BIRTH": 1, "PHONE": 1, "ADDRESS": 2,
    "PERSON_NAME": 3, "COMPANY_NAME": 4,
}


def resolve_overlaps(spans):
    """Keep the highest-priority, then longest, span when spans overlap."""
    spans = sorted(spans, key=lambda s: (s[0], -(s[1] - s[0])))
    kept = []
    occupied = []  # list of (start, end)
    # sort candidates by priority then length so good spans claim territory first
    for s in sorted(spans, key=lambda s: (CATEGORY_PRIORITY.get(s[2], 9), -(s[1] - s[0]))):
        start, end = s[0], s[1]
        if any(start < o[1] and end > o[0] for o in occupied):
            continue
        occupied.append((start, end))
        kept.append(s)
    return sorted(kept, key=lambda s: s[0])


def build_gazetteer(all_text, use_generic_fallback=False):
    names = {}
    companies = {}
    for m in det.find_names(all_text, use_generic_fallback=use_generic_fallback):
        names[m[3]] = names.get(m[3], 0) + 1
    for m in det.find_companies(all_text):
        companies[m[3]] = companies.get(m[3], 0) + 1
    return set(names), set(companies)


def gazetteer_spans(text, gaz_names, gaz_companies):
    # Case-insensitive so ALL-CAPS cover-page duplicates of a known mixed-case
    # entity (e.g. "KUSHAL SUBBAYYA HEGDE" vs "Kushal Subbayya Hegde") are
    # also caught, at the cost of always emitting the same-cased fake value.
    spans = []
    for name in gaz_names:
        for m in re.finditer(re.escape(name), text, re.IGNORECASE):
            spans.append((m.start(), m.end(), "PERSON_NAME", name))
    for name in gaz_companies:
        for m in re.finditer(re.escape(name), text, re.IGNORECASE):
            spans.append((m.start(), m.end(), "COMPANY_NAME", name))
    return spans


def detect_all(text, gaz_names, gaz_companies, use_generic_fallback=False):
    spans = []
    for fn in det.DETECTORS:
        if fn is det.find_names:
            spans.extend(fn(text, use_generic_fallback=use_generic_fallback))
        else:
            spans.extend(fn(text))
    spans.extend(gazetteer_spans(text, gaz_names, gaz_companies))
    return resolve_overlaps(spans)


def apply_redactions_to_runs(run_bounds, spans, factory, audit_rows):
    changed = False
    spans = sorted(spans, key=lambda s: s[0], reverse=True)
    new_run_texts = [r.text for _, _, r in run_bounds]
    
    for start, end, category, original in spans:
        fake = factory.get(category, original)
        audit_rows.append((category, original, fake))
        changed = True
        
        first_overlap_idx = -1
        for i in range(len(run_bounds)):
            r_start, r_end, r = run_bounds[i]
            if start < r_end and end > r_start:
                first_overlap_idx = i
                break
                
        if first_overlap_idx == -1:
            continue
            
        r_start, r_end, _ = run_bounds[first_overlap_idx]
        local_start = max(0, start - r_start)
        local_end = min(r_end - r_start, end - r_start)
        
        text = new_run_texts[first_overlap_idx]
        new_run_texts[first_overlap_idx] = text[:local_start] + fake + text[local_end:]
        
        for i in range(first_overlap_idx + 1, len(run_bounds)):
            r_start, r_end, _ = run_bounds[i]
            if start < r_end and end > r_start:
                local_start = max(0, start - r_start)
                local_end = min(r_end - r_start, end - r_start)
                text = new_run_texts[i]
                new_run_texts[i] = text[:local_start] + "" + text[local_end:]
                
    for i, (_, _, r) in enumerate(run_bounds):
        r.text = new_run_texts[i]
        
    return changed


def iter_paragraphs(document):
    """Yield every paragraph in body, tables (incl. nested), headers, footers."""
    def from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    yield from from_tables(cell.tables)

    yield from document.paragraphs
    yield from from_tables(document.tables)
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            yield from from_tables(part.tables)


def paragraph_text(p):
    return "".join(r.text for r in p.runs) if p.runs else p.text


def redact_docx(in_path, out_path, audit_path=None):
    document = Document(in_path)

    # ---- Pass 1: build gazetteer -------------------------------------
    all_text_chunks = [paragraph_text(p) for p in iter_paragraphs(document)]
    gaz_names, gaz_companies = build_gazetteer("\n".join(all_text_chunks))

    # ---- Pass 2: redact -----------------------------------------------
    factory = FakeValueFactory()
    audit_rows = []
    n_changed = 0

    def redact_paragraph(p):
        nonlocal n_changed
        original = paragraph_text(p)
        if not original.strip():
            return
        spans = detect_all(original, gaz_names, gaz_companies, use_generic_fallback=False)
        if spans:
            run_bounds = []
            cursor = 0
            for r in p.runs:
                length = len(r.text)
                run_bounds.append((cursor, cursor + length, r))
                cursor += length
            if apply_redactions_to_runs(run_bounds, spans, factory, audit_rows):
                n_changed += 1

    def redact_cell(cell):
        """
        Table cells can split a single sentence across multiple internal
        paragraphs. To preserve line breaks and formatting, we map runs across 
        all paragraphs and apply the spans using run bounds.
        """
        nonlocal n_changed
        paras = [p for p in cell.paragraphs if p.runs]
        if len(paras) <= 1:
            for p in cell.paragraphs:
                redact_paragraph(p)
            return
            
        texts = [paragraph_text(p) for p in paras]
        joined = "\n".join(texts)
        if not joined.strip():
            return
            
        spans = detect_all(joined, gaz_names, gaz_companies, use_generic_fallback=False)
        if not spans:
            return
            
        run_bounds = []
        cursor = 0
        for i, p in enumerate(paras):
            for r in p.runs:
                length = len(r.text)
                run_bounds.append((cursor, cursor + length, r))
                cursor += length
            if i < len(paras) - 1:
                cursor += 1  # For the injected \n
                
        if apply_redactions_to_runs(run_bounds, spans, factory, audit_rows):
            n_changed += 1

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    redact_cell(cell)
                    walk_tables(cell.tables)

    for p in document.paragraphs:
        redact_paragraph(p)
    walk_tables(document.tables)
    for section in document.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                redact_paragraph(p)
            walk_tables(part.tables)

    document.save(out_path)

    if audit_path:
        with open(audit_path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["category", "original", "replacement"])
            for row in audit_rows:
                w.writerow(row)

    return {
        "paragraphs_changed": n_changed,
        "total_redactions": len(audit_rows),
        "by_category": _tally(audit_rows),
        "gazetteer_names": sorted(gaz_names),
        "gazetteer_companies": sorted(gaz_companies),
    }


def _tally(audit_rows):
    tally = {}
    for category, _, _ in audit_rows:
        tally[category] = tally.get(category, 0) + 1
    return tally


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 redactor.py <input.docx> <output.docx> [audit.csv]")
        sys.exit(1)
    in_path, out_path = sys.argv[1], sys.argv[2]
    audit_path = sys.argv[3] if len(sys.argv) > 3 else None
    stats = redact_docx(in_path, out_path, audit_path)
    print(f"Paragraphs changed: {stats['paragraphs_changed']}")
    print(f"Total redactions:  {stats['total_redactions']}")
    for cat, count in sorted(stats["by_category"].items()):
        print(f"  {cat:15s} {count}")
